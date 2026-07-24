"""
File Parser Utility
===================
Extracts raw text and structured fields from PDF, DOCX, and TXT files.
"""

import io
import re
from pathlib import Path
from typing import Optional

from loguru import logger

from app.schemas.schemas import ParsedResumeData


class FileParser:
    """
    Parses uploaded resume files (PDF / DOCX / TXT) into structured data.

    Usage:
        parser = FileParser()
        data = parser.parse(file_bytes, "resume.pdf")
    """

    # ── Regex patterns ─────────────────────────────────────────
    EMAIL_RE = re.compile(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}")
    PHONE_RE = re.compile(
        r"(\+?\d{1,3}[\s\-]?)?\(?\d{3}\)?[\s\-]?\d{3}[\s\-]?\d{4}"
    )
    URL_RE = re.compile(
        r"(https?://[^\s]+|linkedin\.com/[^\s]+|github\.com/[^\s]+)"
    )

    SECTION_HEADERS = {
        "education": ["education", "academic", "qualification", "degree"],
        "experience": ["experience", "work history", "employment", "professional"],
        "skills": ["skills", "technical skills", "core competencies", "technologies"],
        "projects": ["projects", "personal projects", "key projects"],
        "certifications": ["certifications", "certificates", "licenses"],
        "achievements": ["achievements", "awards", "honors", "accomplishments"],
        "summary": ["summary", "objective", "profile", "about me"],
    }

    def parse(self, file_bytes: bytes, filename: str) -> ParsedResumeData:
        """
        Main entry point. Detects file type from extension and delegates.

        Args:
            file_bytes: Raw bytes of the uploaded file.
            filename: Original filename including extension.

        Returns:
            ParsedResumeData with all extracted fields.
        """
        ext = Path(filename).suffix.lower().strip(".")
        if ext == "pdf":
            raw_text = self._parse_pdf(file_bytes)
        elif ext == "docx":
            raw_text = self._parse_docx(file_bytes)
        elif ext == "txt":
            raw_text = file_bytes.decode("utf-8", errors="ignore")
        else:
            raise ValueError(f"Unsupported file type: {ext}")

        return self._extract_fields(raw_text)

    # ── Private parsers ────────────────────────────────────────
    def _parse_pdf(self, file_bytes: bytes) -> str:
        """Extract text from PDF using pdfplumber (more reliable than PyMuPDF for text)."""
        try:
            import pdfplumber
            with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
                pages = []
                for page in pdf.pages:
                    text = page.extract_text()
                    if text:
                        pages.append(text)
            return "\n".join(pages)
        except Exception as e:
            logger.warning(f"pdfplumber failed, falling back to PyMuPDF: {e}")
            return self._parse_pdf_fitz(file_bytes)

    def _parse_pdf_fitz(self, file_bytes: bytes) -> str:
        """Fallback PDF parser using PyMuPDF (fitz)."""
        import fitz  # PyMuPDF
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text

    def _parse_docx(self, file_bytes: bytes) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)

    # ── Field extraction ───────────────────────────────────────
    def _extract_fields(self, raw_text: str) -> ParsedResumeData:
        """Parse raw text into structured resume fields."""
        lines = [line.strip() for line in raw_text.split("\n") if line.strip()]

        return ParsedResumeData(
            name=self._extract_name(lines),
            email=self._extract_email(raw_text),
            phone=self._extract_phone(raw_text),
            links=self._extract_links(raw_text),
            summary=self._extract_section(lines, "summary"),
            education=self._extract_section_list(lines, "education"),
            experience=self._extract_section_list(lines, "experience"),
            skills=self._extract_skills(lines),
            projects=self._extract_section_list(lines, "projects"),
            certifications=self._extract_section_list(lines, "certifications"),
            achievements=self._extract_section_list(lines, "achievements"),
            raw_text=raw_text,
        )

    def _extract_name(self, lines: list[str]) -> Optional[str]:
        """Heuristic: first non-empty line that looks like a name."""
        for line in lines[:5]:
            # Skip lines that look like contact info
            if self.EMAIL_RE.search(line) or self.PHONE_RE.search(line):
                continue
            # Skip lines that are section headers
            if any(kw in line.lower() for kws in self.SECTION_HEADERS.values() for kw in kws):
                continue
            # A name is typically 2–4 words, title-cased, no special chars
            words = line.split()
            if 2 <= len(words) <= 4 and all(w.replace("-", "").isalpha() for w in words):
                return line
        return None

    def _extract_email(self, text: str) -> Optional[str]:
        match = self.EMAIL_RE.search(text)
        return match.group() if match else None

    def _extract_phone(self, text: str) -> Optional[str]:
        match = self.PHONE_RE.search(text)
        return match.group() if match else None

    def _extract_links(self, text: str) -> list[str]:
        return list(set(self.URL_RE.findall(text)))

    def _extract_section(self, lines: list[str], section: str) -> Optional[str]:
        """Extract the first paragraph under a section header."""
        content = self._extract_section_list(lines, section)
        return "\n".join(content) if content else None

    def _extract_section_list(self, lines: list[str], section: str) -> list[str]:
        """Extract all bullet lines under a section header."""
        headers = self.SECTION_HEADERS.get(section, [])
        in_section = False
        content: list[str] = []

        for line in lines:
            lower_line = line.lower()
            # Check if this line is the target section header
            if any(h in lower_line for h in headers) and len(line) < 50:
                in_section = True
                continue
            # Check if we've hit another section header → stop
            if in_section:
                is_other_header = any(
                    any(h in lower_line for h in kws) and len(line) < 50
                    for sec, kws in self.SECTION_HEADERS.items()
                    if sec != section
                )
                if is_other_header:
                    break
                if line:
                    content.append(line)

        return content[:20]  # Cap to avoid noise

    def _extract_skills(self, lines: list[str]) -> list[str]:
        """Extract individual skills from the skills section."""
        raw = self._extract_section_list(lines, "skills")
        skills: list[str] = []
        for line in raw:
            # Split by common delimiters
            parts = re.split(r"[,|•·\u2022\t]+", line)
            skills.extend([p.strip() for p in parts if len(p.strip()) > 1])
        return list(dict.fromkeys(skills))[:50]  # Deduplicate, cap at 50
